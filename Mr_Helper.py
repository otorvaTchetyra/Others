import requests
from PIL import Image, ImageDraw, ImageFont
from docx import Document

# Step 1: Поиск информации в интернете
def fetch_news():
    url = 'https://newsapi.org/v2/everything'
    params = {
        'q': 'early education math reading',  # ключевые слова
        'apiKey': '9e3c2104624d4b8982ba5bb8908aded3',  # ваш API ключ
        'sortBy': 'publishedAt',
        'language': 'ru'
    }
    response = requests.get(url, params=params)
    news = response.json()
    return news['articles'][0]['title'], news['articles'][0]['description']

# Step 2: Создание текста для блога
def create_blog_text(title, description):
    return f"Заголовок: {title}\n\n{description}\n\nРаннее обучение детей — это важный шаг к их успешному будущему!"

# Step 3: Создание текста для изображения
def create_image_text(blog_text):
    return "Раннее обучение — ключ к успешному будущему вашего ребенка!"

# Step 4: Создание изображения
def create_image(text):
    img = Image.new('RGB', (800, 400), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    font = ImageFont.load_default()
    d.text((10, 10), text, fill=(0, 0, 0), font=font)
    img.save('output_image.png')

# Step 5: Сохранение результатов
def save_blog_text(blog_text):
    doc = Document()
    doc.add_heading('Блог о раннем обучении', level=1)
    doc.add_paragraph(blog_text)
    doc.save('blog_text.docx')

def save_image_text(image_text):
    with open('image_text.txt', 'w') as f:
        f.write(image_text)

# Основной поток программы
if __name__ == "__main__":
    title, description = fetch_news()
    blog_text = create_blog_text(title, description)
    image_text = create_image_text(blog_text)
    
    save_blog_text(blog_text)
    create_image(image_text)
    save_image_text(image_text)

    print("Все результаты сохранены.")
